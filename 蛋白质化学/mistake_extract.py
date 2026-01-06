import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_docx_content(docx_path):
    """读取.docx文件的内容"""
    try:
        # 直接使用python-docx库读取
        doc = Document(docx_path)
        content = []
        for para in doc.paragraphs:
            if para.text.strip():  # 只添加非空段落
                content.append(para.text)
        return '\n'.join(content)
    except Exception as e:
        print(f"读取文档失败: {e}")
        return ""

def extract_questions(content, question_numbers):
    """提取指定题号的题目"""
    questions = {}
    current_question = None
    current_content = []
    
    lines = content.split('\n')
    
    for line in lines:
        # 检查是否是题目的开始（以数字开头，可能带有特殊字符）
        match = re.match(r'^(\d+)[、\.]', line.strip())
        if match:
            # 保存上一个题目
            if current_question is not None:
                questions[current_question] = '\n'.join(current_content)
            
            # 开始新题目
            question_num = int(match.group(1))
            current_question = question_num
            current_content = [line]
        elif current_question is not None:
            # 继续当前题目
            current_content.append(line)
    
    # 保存最后一个题目
    if current_question is not None:
        questions[current_question] = '\n'.join(current_content)
    
    # 提取指定题号
    selected_questions = {}
    for num in question_numbers:
        if num in questions:
            selected_questions[num] = questions[num]
        else:
            print(f"警告：未找到第 {num} 题")
    
    return selected_questions

def create_word_document(questions, output_filename):
    """创建Word文档"""
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('选择题提取结果', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加说明
    doc.add_paragraph(f'提取了以下题号：{", ".join(map(str, sorted(questions.keys())))}')
    doc.add_paragraph()
    
    # 按题号顺序添加题目
    for num in sorted(questions.keys()):
        # 添加题目
        question_text = questions[num]
        # 处理题目文本中的特殊格式
        lines = question_text.split('\n')
        
        for i, line in enumerate(lines):
            if i == 0:  # 第一行是题号，加粗
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            else:
                if line.strip():  # 非空行
                    doc.add_paragraph(line)
        
        # 添加空行分隔
        doc.add_paragraph()
    
    # 保存文档
    doc.save(output_filename)
    print(f"文档已保存为：{output_filename}")

def main():
    # 读取文档内容
    docx_path = '选择题.docx'
    content = read_docx_content(docx_path)
    
    if not content:
        print("无法读取文档内容，请确保文件格式正确")
        return
    
    # 获取用户输入的题号
    input_str = input("请输入要提取的题号（用逗号隔开，例如：1,3,5,7）：")
    
    try:
        # 解析输入
        question_numbers = [int(num.strip()) for num in input_str.split(',')]
        
        # 提取题目
        selected_questions = extract_questions(content, question_numbers)
        
        if selected_questions:
            # 生成输出文件名
            numbers_str = '_'.join(map(str, sorted(question_numbers)))
            output_filename = f"提取题目_{numbers_str}.docx"
            
            # 创建Word文档
            create_word_document(selected_questions, output_filename)
            
            print(f"成功提取了 {len(selected_questions)} 个题目")
        else:
            print("未找到任何匹配的题目")
            
    except ValueError:
        print("输入格式错误，请确保输入的是用逗号隔开的数字")
    except Exception as e:
        print(f"发生错误：{e}")

if __name__ == "__main__":
    main()